from fastapi import FastAPI, UploadFile, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import os
import sys
import markdown
from xhtml2pdf import pisa
from docx import Document
from io import BytesIO
from fastapi.responses import FileResponse, StreamingResponse

# Add project root to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from src.orchestrator import Orchestrator
from dotenv import load_dotenv
from src.utils.report_formatter import ReportFormatter

# Load env variables
load_dotenv()

app = FastAPI(title="Multi-Agent Researcher API")

# Mount static files
# Mount static files
app.mount("/assets", StaticFiles(directory="frontend/dist/assets"), name="assets")

@app.get("/")
async def read_index():
    return FileResponse('frontend/dist/index.html')

class ResearchRequest(BaseModel):
    topic: str
    custom_prompt: str = None

@app.post("/api/research")
async def conduct_research(request: ResearchRequest):
    """
    Endpoint to trigger the research process.
    """
    if not os.getenv("GOOGLE_API_KEY"):
        return {"error": "GOOGLE_API_KEY not found. Please set it in .env file."}

    topic = request.topic
    custom_prompt = request.custom_prompt
    orchestrator = Orchestrator()
    
    # Run research (synchronously for now, but could be async or background task)
    # Since orchestrator.run is blocking and time-consuming, ideally we'd use background tasks
    # or a proper job queue. For simplicity here, we'll wait.
    try:
        report = orchestrator.run(topic, custom_prompt)
        return {"report": report}
    except Exception as e:
        return {"error": str(e)}

# --- Multi-Stage Workflow Endpoints ---

class PlanRequest(BaseModel):
    topic: str
    custom_prompt: str = None

class SubTopicInstruction(BaseModel):
    topic: str
    instructions: str = None

class ResearchPhaseRequest(BaseModel):
    sub_topics: list[SubTopicInstruction] = []

class SummarizeRequest(BaseModel):
    topic: str
    research_findings: dict
    sources: list = []
    custom_prompt: str = None

@app.post("/api/plan")
async def plan_research(request: PlanRequest):
    """Stage 1: Generate plan"""
    if not os.getenv("GOOGLE_API_KEY"):
         return {"error": "GOOGLE_API_KEY not found."}
    
    orchestrator = Orchestrator()
    try:
        sub_topics = orchestrator.plan_research(request.topic, request.custom_prompt)
        return {"sub_topics": sub_topics}
    except Exception as e:
        return {"error": str(e)}

@app.post("/api/research_phase")
async def execute_research(request: ResearchPhaseRequest):
    """Stage 2: Execute research on confirmed sub-topics"""
    orchestrator = Orchestrator()
    try:
        findings, sources = orchestrator.execute_research(request.sub_topics)
        # Return structured findings for frontend editing
        return {
            "findings": findings, 
            "sources": sources
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/api/summarize")
async def generate_summary(request: SummarizeRequest):
    """Stage 3: Generate final report from confirmed findings"""
    orchestrator = Orchestrator()
    try:
        report = orchestrator.generate_summary(request.topic, request.research_findings, request.sources, custom_prompt=request.custom_prompt)
        return {"report": report}
    except Exception as e:
        return {"error": str(e)}

class ExportRequest(BaseModel):
    content: str

@app.post("/api/export/pdf")
async def export_pdf(request: ExportRequest):
    try:
        formatter = ReportFormatter()
        pdf_buffer = formatter.generate_pdf(request.content)
        
        return StreamingResponse(
            pdf_buffer, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=report.pdf"}
        )
    except Exception as e:
        return {"error": str(e)}

@app.post("/api/export/docx")
async def export_docx(request: ExportRequest):
    document = Document()
    
    # Simple conversion: Split by lines and add paragraphs
    # For a robust solution, we'd parse the markdown properly. 
    # Here we'll do a basic pass:
    # - Lines starting with # become Heading 1
    # - Lines starting with ## become Heading 2
    # - Others are paragraphs
    
    for line in request.content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
             document.add_paragraph(line[2:], style='List Bullet')
        else:
            document.add_paragraph(line)
            
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=report.docx"}
    )

if __name__ == "__main__":
    import uvicorn
    # Clean up any previous runs if port is stuck (not possible from here but good practice)
    uvicorn.run(app, host="127.0.0.1", port=8000)
